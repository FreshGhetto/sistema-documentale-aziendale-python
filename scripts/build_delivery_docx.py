from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent.parent
SOURCE = BASE_DIR / "docs" / "documento_consegna.md"
TARGET = BASE_DIR / "docs" / "documento_consegna.docx"


def add_markdown_line(document: Document, raw_line: str) -> None:
    line = raw_line.strip()
    if not line:
        document.add_paragraph("")
        return
    if line.startswith("# "):
        document.add_heading(line[2:], level=1)
        return
    if line.startswith("## "):
        document.add_heading(line[3:], level=2)
        return
    if line.startswith("### "):
        document.add_heading(line[4:], level=3)
        return
    if line.startswith("- "):
        document.add_paragraph(line[2:], style="List Bullet")
        return
    if line.startswith("```"):
        return
    document.add_paragraph(line)


def build_document() -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for raw_line in SOURCE.read_text(encoding="utf-8").splitlines():
        add_markdown_line(document, raw_line)

    document.save(TARGET)


if __name__ == "__main__":
    build_document()
