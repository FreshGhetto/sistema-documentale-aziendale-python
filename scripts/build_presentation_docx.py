from pathlib import Path

from docx import Document
from docx.shared import Pt


BASE_DIR = Path(__file__).resolve().parent.parent
SOURCE = BASE_DIR / "docs" / "presentazione_demo.md"
TARGET = BASE_DIR / "docs" / "presentazione_demo.docx"


def build_document() -> None:
    document = Document()
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            document.add_paragraph("")
            continue
        if line.startswith("# "):
            document.add_heading(line[2:], level=1)
            continue
        if line.startswith("## "):
            document.add_heading(line[3:], level=2)
            continue
        if line.startswith("- "):
            document.add_paragraph(line[2:], style="List Bullet")
            continue
        if line.startswith("1. ") or line.startswith("2. ") or line.startswith("3. ") or line.startswith("4. ") or line.startswith("5. "):
            document.add_paragraph(line[3:], style="List Number")
            continue
        document.add_paragraph(line)

    document.save(TARGET)


if __name__ == "__main__":
    build_document()
